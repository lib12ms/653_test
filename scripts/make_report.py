"""오늘 작업 내용을 Word 문서로 생성."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import datetime

doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)

def set_font(run, bold=False, size=10, color=None):
    run.font.name = '맑은 고딕'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_font(run, bold=True, size=14, color=(0x1F, 0x49, 0x7D))
    elif level == 2:
        set_font(run, bold=True, size=12, color=(0x2E, 0x74, 0xB5))
    elif level == 3:
        set_font(run, bold=True, size=11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10)
    return p

def add_bullet(text, indent=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    set_font(run, size=10)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # 헤더
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, size=9)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        # 헤더 배경색
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'BDD7EE')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    # 데이터
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r + 1, c)
            cell.text = val
            for run in cell.paragraphs[0].runs:
                set_font(run, size=9)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

# ── 제목 ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('653 자유주제어 자동생성 시스템\n품질검증 및 프롬프트 고도화 작업일지')
set_font(run, bold=True, size=16, color=(0x1F, 0x49, 0x7D))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run(f'작업일자: 2026-05-23')
set_font(run, size=10, color=(0x70, 0x70, 0x70))

doc.add_paragraph()

# ── 1. 테스트 개요 ───────────────────────────────────────────────────────────
add_heading('1. 테스트 개요', level=1)

add_heading('목적', level=3)
add_body('MARC 653 자유주제어 자동생성 시스템의 품질 검증 및 프롬프트 고도화')
add_body('테스트 → 결과 CSV 검토 → 문제 키워드 확인 → 프롬프트/필터 수정의 반복 사이클로 진행')

add_heading('데이터 선정 기준', level=3)
add_bullet('출판일 기준 최근 1개월 이내(2026-04-23 이후) 신간만 대상')
add_bullet('출판일 미기재 또는 파싱 불가 도서는 제외')
add_bullet('알라딘 TTB API ItemNewAll 쿼리로 분야별 최신순 수집')
add_bullet('분야별 알라딘 카테고리 키워드 필터 적용 → 타 분야 유입 방지')

add_heading('신간 도서를 기준으로 삼은 이유', level=3)
add_bullet('현재 이용자가 실제로 검색하는 최신 도서에 시스템이 제대로 작동하는지 검증하기 위함')
add_bullet('신간은 메타데이터(소개문·목차)가 부실한 경우가 많아 프롬프트의 한계를 조기에 발견할 수 있음')
add_bullet('도서관 현장에서 653 필드 자동생성이 필요한 시점이 바로 신간 입수 단계임')

add_heading('테스트 실행 결과', level=3)
add_table(
    ['회차', '일시', '분야당 목표', '총 권수', '성공률', '비고'],
    [
        ['1차', '2026-05-23 20:09', '15권', '148권', '-', '종교 분야 별도 재실행 후 병합'],
        ['2차', '2026-05-23 22:04', '10권', '96권', '100%', '1차 ISBN 전량 제외, 신규 도서만 대상'],
    ]
)
add_body('※ 종교 분야: 신간 부족으로 6권만 수집')
add_body('※ 2차 테스트는 --exclude 옵션으로 1차 148개 ISBN 자동 제외')

# ── 2. 프롬프트 수정 — 전 분야 공통 ─────────────────────────────────────────
add_heading('2. 프롬프트 수정 — 전 분야 공통 (_STATIC_INSTRUCTIONS)', level=1)

add_table(
    ['항목', '변경 내용'],
    [
        ['구체성 원칙', '금지 분류명에 과학, 인문학, 자기계발, 예술, 교육 추가 명시'],
        ['유사어 중복 규칙', '~적 형용사형과 명사형이 겹치면 명사형 선택 기준 추가'],
        ['배경 키워드 규칙', '시대·지역 배경은 설명·목차 명시 근거 있을 때만 추출, 배경지식 추정 금지'],
        ['국가+장르 규칙', '한국·국내+장르만 차단, 외국 국가+장르(일본소설, 러시아문학 등) 허용으로 완화'],
    ]
)

# ── 3. 프롬프트 수정 — 분야별 ─────────────────────────────────────────────
add_heading('3. 프롬프트 수정 — 분야별', level=1)

add_heading('문학', level=2)
add_bullet('추상 감상어 목록 명시 금지 (운명, 인연, 만남, 기억, 시간, 사랑 등)')
add_bullet('[배경 키워드 필수 발굴] → [배경 키워드 발굴 — 근거 필수]로 변경')
add_bullet('소개문 빈약 시 대안 발굴 3단계 추가:')
add_bullet('① 분류 꼬리 하위 장르명', indent=2)
add_bullet('② 제목 함의 독자 상황', indent=2)
add_bullet('③ 한 단계 구체화된 문학 형식 (단순 에세이·소설 금지, 철학에세이·전쟁소설 등 허용)', indent=2)

add_heading('에세이', level=2)
add_bullet('동일 추상 감상어·단독 명사 목록 명시 금지')
add_bullet('소개문 빈약 시 동일 3단계 대안 발굴 추가')

add_heading('인문학', level=2)
add_bullet('[분류명 금지] 항목 추가: 인문학, 철학, 역사 단독 사용 금지')
add_bullet('철학자·사상가 처리 원칙 재정비:')
add_bullet('① 제목·저자에 있는 사상가명 → 단독 절대 금지, 사상가명+분야 필수', indent=2)
add_bullet('② 설명·목차에만 있는 사상가명 → 인명 단독만 허용, 분야 추정 결합 금지', indent=2)

# ── 4. 후처리 필터 수정 ────────────────────────────────────────────────────
add_heading('4. 후처리 필터 수정 (ai_service.py)', level=1)

add_table(
    ['항목', '변경 내용'],
    [
        ['LOW_VALUE_KEYWORDS', '자기계발 추가'],
        ['_LOW_VALUE_SUFFIX_RE', '탐구, 통찰, 담론, 해석 접미어 추가'],
        ['_LIT_ABSTRACT_NOUNS (신규)', '문학·에세이 전용 단독 추상명사 집합 생성 및 필터 연결'],
        ['_COUNTRY_GENRE_RE', '한국·국내+장르만 차단하도록 축소'],
        ['_EXTENDED_COUNTRY_GENRE_RE', '동일'],
    ]
)

# ── 5. 테스트 스크립트 수정 ────────────────────────────────────────────────
add_heading('5. 테스트 스크립트 수정 (test_new_653.py)', level=1)

add_bullet('--exclude 옵션 추가: 이전 테스트 CSV의 ISBN을 자동 제외')
add_bullet('per_cid 계산 로직 개선: exclude 크기 반영하여 수집 부족 문제 해결')

# ── 6. 보류 항목 ────────────────────────────────────────────────────────────
add_heading('6. 보류 항목', level=1)

add_bullet('인문교양 류 분류 라벨 처리')
add_bullet('교육·좋은부모 분야 교육법 키워드(유대인교육법 류) 처리')
add_bullet('기타 분야 여행·전집 CID (수동 테스트로 대체)')

# 저장
out = Path('d:/653_test') / f'653_작업일지_20260523.docx'
doc.save(out)
print(f'saved: {out}')
